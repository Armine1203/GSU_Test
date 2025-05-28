
from django.contrib.contenttypes.models import ContentType
from django.db import transaction
from django.views.decorators.http import require_POST
from django.db.models import Q
from .models import TestQuestion, MidtermExam, ExamResult, StudentAnswer, Mark, QuestionCategory
from django.db import models
from django.db.models import Prefetch, Avg, Count, F, Sum, Exists, OuterRef, When, Case, BooleanField
from django.shortcuts import render, redirect
from django.urls import reverse
from django.views import View
from .forms import SubjectForm, MidtermExamForm
from .models import TestQuestion, Lecturer, Subject, Group, StudentAnswer, ExamResult, Mark, MidtermExam, Student, \
    LiveStudentExam
from django.contrib import messages
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from django.http import JsonResponse, HttpResponseForbidden, HttpResponseNotAllowed
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from django.views.generic import View
from docx import Document
from django.views.decorators.csrf import csrf_exempt
from .forms import FeedbackForm
from django.views.generic import View
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator


# Create your views here.
@method_decorator(login_required, name="dispatch")
class Quiz(View):
    def get(self, request):
        exam_id = request.GET.get('exam_id')

        if exam_id:
            try:
                midterm_exam = get_object_or_404(MidtermExam, id=exam_id)

                # Check if student already attempted this exam
                if Mark.objects.filter(user=request.user, exam=midterm_exam).exists():
                    messages.error(request, "Դուք արդեն անցել եք այս քննությունը")
                    return redirect("student_dashboard")

                # Try to get LiveStudentExam if it exists
                try:
                    live_exam = LiveStudentExam.objects.get(
                        exam=midterm_exam,
                        student=request.user.student
                    )
                    questions = live_exam.questions.all()
                    time_limit = live_exam.time_limit
                except LiveStudentExam.DoesNotExist:
                    # Fall back to MidtermExam questions if no LiveStudentExam exists
                    questions = midterm_exam.questions.all()
                    time_limit = midterm_exam.time_limit

                return render(request, "quiz/quiz.html", {
                    "questions": questions,
                    "time_limit": time_limit * 60,
                    "start_time": timezone.now().isoformat(),
                    "exam": midterm_exam
                })

            except MidtermExam.DoesNotExist:
                messages.error(request, "Քննությունը չի գտնվել")
                return redirect("student_dashboard")

        messages.error(request, "Չի նշվել քննություն")
        return redirect("student_dashboard")

    def post(self, request):
        exam_id = request.POST.get('exam_id')
        try:
            # Always get the MidtermExam instance first
            midterm_exam = MidtermExam.objects.get(id=exam_id)

            # Check if student already took the exam
            if Mark.objects.filter(user=request.user, exam=midterm_exam).exists():
                messages.error(request, "Դուք արդեն անցել եք այս քննությունը")
                return redirect("student_dashboard")

            # Check if LiveStudentExam exists
            try:
                live_exam = LiveStudentExam.objects.get(exam=midterm_exam, student=request.user.student)
                questions = live_exam.questions.all()
                exam_for_result = live_exam
            except LiveStudentExam.DoesNotExist:
                questions = midterm_exam.questions.all()
                exam_for_result = midterm_exam

            correct = 0
            student_answers = []

            for question in questions:
                user_answer = request.POST.get(f'q{question.id}o')
                if not user_answer:
                    continue

                is_correct = user_answer == question.correct_option
                if is_correct:
                    correct += question.score  # ✅ use question's score instead of 1

                answer = StudentAnswer.objects.create(
                    question=question,
                    answer=user_answer,
                    is_correct=is_correct
                )
                student_answers.append(answer)

            # Save total and got score
            mark = Mark.objects.create(
                user=request.user,
                exam=midterm_exam,  # ✅ Always save MidtermExam here
                total=sum(q.score for q in questions),
                got=correct
            )

            exam_result = ExamResult.objects.create(
                exam=exam_for_result,  # This can be LiveStudentExam or MidtermExam
                student=request.user.student,
                score=correct
            )
            exam_result.answers.set(student_answers)

            return redirect("result_detail", result_id=mark.id)

        except MidtermExam.DoesNotExist:
            messages.error(request, "Քննությունը չի գտնվել")
            return redirect("student_dashboard")


@method_decorator(login_required, name="dispatch")
class AddQuestion(View):
    def get(self, request):
        if not hasattr(request.user, 'lecturer'):
            messages.error(request, "Միայն դասախոսները կարող են հարցեր ավելացնել")
            return redirect("index")

        # Get parameters with defaults
        try:
            question_count = int(request.GET.get('count', 2))
            question_count = max(1, min(100, question_count))  # Limit between 1-100
        except ValueError:
            question_count = 2

        subject_id = request.GET.get('subject')
        lecturer = request.user.lecturer
        subject = None

        if subject_id:
            try:
                subject = Subject.objects.get(id=subject_id)
                if subject not in lecturer.subjects.all():
                    subject = None
            except Subject.DoesNotExist:
                pass

        return render(
            request,
            "quiz/add_questions.html",
            {
                "questions": range(1, question_count + 1),
                "lecturer": lecturer,
                "subjects": lecturer.subjects.all(),
                "selected_subject": subject,
                "question_count": question_count,
                "categories": QuestionCategory.objects.filter(created_by=request.user),  # Fixed this line
            }
        )

    def post(self, request):
        if not hasattr(request.user, 'lecturer'):
            messages.error(request, "Միայն դասախոսները կարող են հարցեր ավելացնել")
            return redirect("index")

        lecturer = request.user.lecturer
        subject_id = request.POST.get("subject")

        try:
            subject = Subject.objects.get(id=subject_id)
            if subject not in lecturer.subjects.all():
                messages.error(request, "You don't teach this subject")
                return redirect("add_question")
        except Subject.DoesNotExist:
            messages.error(request, "Invalid subject selected")
            return redirect("add_question")

        count, already_exists = 0, 0
        for i in range(1, int(request.POST.get("question_count", 10)) + 1):
            data = request.POST
            files = request.FILES

            q = data.get(f"q{i}", "").strip()
            if not q:  # Skip empty questions
                continue

            o1 = data.get(f"q{i}o1", "").strip()
            o2 = data.get(f"q{i}o2", "").strip()
            o3 = data.get(f"q{i}o3", "").strip()
            o4 = data.get(f"q{i}o4", "").strip()
            co = data.get(f"q{i}c", "")

            score = int(data.get(f"q{i}score", 1))
            image = files.get(f"q{i}image")

            if TestQuestion.objects.filter(question=q, subject=subject).exists():
                already_exists += 1
                continue

            category_id = data.get(f"q{i}category")
            category = None
            if category_id:
                try:
                    category = QuestionCategory.objects.get(id=category_id, created_by=request.user)  # Fixed this line
                except QuestionCategory.DoesNotExist:
                    pass

            question = TestQuestion(
                subject=subject,
                question=q,
                option1=o1,
                option2=o2,
                option3=o3,
                option4=o4,
                correct_option=co,
                score=score,
                creator=request.user,
                image=image,
                category=category,
            )
            question.save()
            count += 1

        if already_exists:
            messages.warning(request, f"{already_exists} questions already exist for this subject")
        if count:
            messages.success(request, f"Հաջողությամբ {count} հարց ավելացվել է {subject.name} առարկայի համար")
        else:
            messages.warning(request, "Հարցեր չկան ավելացվելու համար")

        return redirect("lecturer_dashboard")


@method_decorator(login_required, name="dispatch")
class Result(View):
    def get(self, request):
        results = ExamResult.objects.filter(student__user=request.user)

        # Calculate statistics
        if results.exists():
            total_tests = results.count()
            total_questions = sum(r.total for r in results)
            correct_answers = sum(r.got for r in results)
            avg_percentage = round((correct_answers / total_questions) * 100, 2) if total_questions > 0 else 0
            best_score = results.order_by('-got').first()

            # Add percentage to each result
            for result in results:
                result.percentage = round((result.got / result.total) * 100, 2) if result.total > 0 else 0
        else:
            avg_percentage = 0
            best_score = None

        return render(request, "quiz/results.html", {
            "results": results,
            "avg_percentage": avg_percentage,
            "best_score": best_score
        })


from django.views import View
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.shortcuts import  get_object_or_404
from django.http import HttpResponse
from django.contrib.contenttypes.models import ContentType

from docx import Document
from openpyxl import Workbook

from django.views import View
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from django.db.models import Prefetch
from django.contrib.contenttypes.models import ContentType

from docx import Document
from openpyxl import Workbook

from .models import (
    Student, Lecturer, Group, MidtermExam,
    LiveStudentExam, ExamResult, Subject
)


@method_decorator(login_required, name="dispatch")
class Leaderboard(View):
    def get(self, request, exam_id=None):
        try:
            context = {}

            if hasattr(request.user, 'student'):
                context.update(self._student_view(request))
            elif hasattr(request.user, 'lecturer'):
                context.update(self._lecturer_view(request))
            else:
                return render(request, 'quiz/leaderboard.html', {'error': "Invalid user type"})

            if not context.get('group_exams'):
                context['no_exams'] = True

            return render(request, 'quiz/leaderboard.html', context)

        except Exception as e:
            return render(request, 'quiz/leaderboard.html', {'error': str(e)})

    def _student_view(self, request):
        student = request.user.student
        current_group = student.group

        # Prefetch all related data
        group_exams = MidtermExam.objects.filter(
            group=current_group
        ).select_related('subject').prefetch_related('questions')

        # Get all exam results (both midterm and live exams)
        exam_results = ExamResult.objects.filter(
            student__group=current_group
        ).select_related(
            'student', 'content_type'
        ).order_by('-completed_at')

        # Identify students without results
        students_with_results = set(
            exam_results.values_list('student_id', flat=True).distinct()
        )
        missing_result_students = Student.objects.filter(
            group=current_group
        ).exclude(
            id__in=students_with_results
        )

        return {
            'user_type': 'student',
            'current_group': current_group,
            'group_exams': group_exams,
            'exam_results': exam_results,
            'missing_result_students': missing_result_students,
        }

    def _lecturer_view(self, request):
        lecturer = request.user.lecturer
        subjects = lecturer.subjects.all()
        groups = Group.objects.filter(
            major__in=subjects.values('major')
        ).distinct().select_related('major')

        selected_group_id = request.GET.get('group')
        export_format = request.GET.get('export')

        context = {
            'user_type': 'lecturer',
            'subjects': subjects,
            'groups': groups,
        }

        if selected_group_id:
            selected_group = get_object_or_404(Group, id=selected_group_id)
            group_exams = MidtermExam.objects.filter(
                group=selected_group,
                subject__in=subjects
            ).select_related('subject').prefetch_related('questions')

            exam_results = ExamResult.objects.filter(
                student__group=selected_group,
                content_type__model__in=['midtermexam', 'livestudentexam']
            ).select_related(
                'student', 'content_type'
            ).order_by('-completed_at')

            # Identify students without results
            students_with_results = set(
                exam_results.values_list('student_id', flat=True).distinct()
            )
            missing_result_students = Student.objects.filter(
                group=selected_group
            ).exclude(
                id__in=students_with_results
            )

            context.update({
                'selected_group': selected_group,
                'group_exams': group_exams,
                'exam_results': exam_results,
                'missing_result_students': missing_result_students,
            })

            if export_format in ['word', 'excel']:
                return self.export_results(selected_group, group_exams, exam_results, export_format)

        return context

    def export_results(self, group, exams, results, format):
        if format == 'word':
            return self._export_to_word(group, exams, results)
        elif format == 'excel':
            return self._export_to_excel(group, exams, results)

    def _export_to_word(self, group, exams, results):
        document = Document()
        document.add_heading(f'Քննությունների արդյունքներ - խումբ {group.name}', level=1)

        for exam in exams:
            document.add_heading(f'Առարկա - {exam.subject.name}', level=2)
            document.add_paragraph(f'Ամսաթիվ - {exam.due_date.strftime("%d.%m.%Y %H:%M")}')

            # Create table
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Ուսանող'
            hdr_cells[2].text = 'ID'
            hdr_cells[3].text = 'Միավոր'
            hdr_cells[4].text = 'Տոկոս'

            # Filter and sort results for this exam
            exam_results = [
                r for r in results
                if (r.content_type.model == 'midtermexam' and r.object_id == exam.id) or
                   (r.content_type.model == 'livestudentexam' and r.exam.exam.id == exam.id)
            ]
            exam_results.sort(key=lambda x: x.score, reverse=True)

            # Add data rows
            for idx, result in enumerate(exam_results, start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = f"{result.student.name} {result.student.last_name}"
                row_cells[2].text = result.student.id_number
                row_cells[3].text = f"{result.score}/{result.exam.questions.count()}"
                row_cells[4].text = f"{result.percentage:.1f}%"

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{group.name}_քննության_արդյունքներ.docx"'
        document.save(response)
        return response

    def _export_to_excel(self, group, exams, results):
        wb = Workbook()
        ws = wb.active
        ws.title = f"{group.name} Results"

        # Header row
        ws.append(["#", "Ուսանող", "ID", "Առարկա", "Միավոր", "Տոկոս", "Ամսաթիվ"])

        # Process all exams
        for exam in exams:
            # Filter and sort results for this exam
            exam_results = [
                r for r in results
                if (r.content_type.model == 'midtermexam' and r.object_id == exam.id) or
                   (r.content_type.model == 'livestudentexam' and r.exam.exam.id == exam.id)
            ]
            exam_results.sort(key=lambda x: x.score, reverse=True)

            # Add data rows
            for idx, result in enumerate(exam_results, start=1):
                ws.append([
                    idx,
                    f"{result.student.name} {result.student.last_name}",
                    result.student.id_number,
                    exam.subject.name,
                    f"{result.score}/{result.exam.questions.count()}",
                    f"{result.percentage:.1f}%",
                    exam.due_date.strftime("%d.%m.%Y %H:%M")
                ])

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{group.name}_results.xlsx"'
        wb.save(response)
        return response

@method_decorator(login_required, name="dispatch")
class StudentDashboard(View):
    def get(self, request):
        try:
            student = request.user.student  # More direct access
            now = timezone.now()

            # Get all exams for student's group
            exams = MidtermExam.objects.filter(
                group=student.group
            ).annotate(
                is_completed=Exists(
                    Mark.objects.filter(
                        exam=OuterRef('pk'),
                        user=request.user
                    )
                ),
                is_available=Case(
                    When(due_date__gte=now, then=True),
                    default=False,
                    output_field=BooleanField()
                )
            ).order_by('due_date')

            return render(request, "quiz/student_dashboard.html", {
                'student': student,
                'exams': exams,
                'now': now,  # Pass current time to template
            })

        except AttributeError:  # Catches missing student profile
            messages.error(request, "Student profile not found")
            return redirect("index")

@method_decorator(login_required, name="dispatch")
class LecturerDashboard(View):
    def get(self, request):
        if not hasattr(request.user, 'lecturer'):
            return HttpResponseForbidden("Only lecturers can access this page")

        lecturer = request.user.lecturer
        now = timezone.now()

        # Get lecturer's subjects and questions
        subjects = lecturer.subjects.all()
        total_questions = TestQuestion.objects.filter(subject__in=subjects).count()

        # Get categories
        categories = QuestionCategory.objects.filter(created_by=request.user)
        print("Categories count:", categories.count())  # Debug output

        # Get exams created by this lecturer with annotations
        created_exams = MidtermExam.objects.filter(
            created_by=request.user
        ).annotate(
            question_count=Count('questions'),
            is_active=Case(
                When(due_date__gt=now, then=True),
                default=False,
                output_field=BooleanField()
            )
        ).select_related('subject', 'group').order_by('-due_date')

        form = MidtermExamForm(user=request.user)

        return render(request, "quiz/lecturer_dashboard.html", {
            'lecturer': lecturer,
            'subjects': subjects,
            'total_questions': total_questions,
            'created_exams': created_exams,
            'now': now,
            'form': form,
            'categories': categories,  # Add this line
        })

# Add to views.py
@method_decorator(login_required, name="dispatch")
class ViewQuestions(View):
    def get(self, request, subject_id):
        try:
            subject = Subject.objects.get(id=subject_id)
            lecturer = request.user.lecturer

            # Verify lecturer teaches this subject
            if subject not in lecturer.subjects.all():
                messages.error(request, "Դուք չեք դասավանդում այս առարկան")
                return redirect("lecturer_dashboard")

            questions = TestQuestion.objects.filter(
                subject=subject,
                creator=request.user
            ).order_by('-id')

            return render(request, "quiz/view_questions.html", {
                'subject': subject,
                'questions': questions
            })

        except Subject.DoesNotExist:
            messages.error(request, "Առարկան չի գտնվել")
            return redirect("lecturer_dashboard")

@method_decorator(login_required, name="dispatch")
class DeleteQuestion(View):
    def post(self, request, pk):
        try:
            question = TestQuestion.objects.get(id=pk)
            if question.creator != request.user:
                messages.error(request, "Դուք կարող եք ջնջել միայն ձեր կողմից ստեղծված հարցերը")
                return redirect("lecturer_dashboard")

            question.delete()
            messages.success(request, "Հարցը հաջողությամբ ջնջվել է ")
            return redirect("view_questions", subject_id=question.subject.id)

        except TestQuestion.DoesNotExist:
            messages.error(request, "Հարցը չի գտնվել")
            return redirect("lecturer_dashboard")

@method_decorator(login_required, name="dispatch")
class AddSubject(View):
    def get(self, request):
        form = SubjectForm()
        return render(request, "quiz/add_subject.html", {'form': form})

    def post(self, request):
        form = SubjectForm(request.POST)
        if form.is_valid():
            subject = form.save()
            subject.lecturers.add(request.user)

            # Also add to Lecturer model if exists
            lecturer, created = Lecturer.objects.get_or_create(user=request.user)
            lecturer.subjects.add(subject)

            messages.success(request, "Առարկան հաջողությամբ ավելացվել է")
            return redirect('lecturer_dashboard')
        return render(request, "quiz/add_subject.html", {'form': form})

@method_decorator(login_required, name="dispatch")
class SelectQuestionCount(View):
    def get(self, request):
        if not hasattr(request.user, 'lecturer'):
            messages.error(request, "Only lecturers can add questions")
            return redirect("index")

        lecturer = request.user.lecturer
        return render(request, "quiz/select_question_count.html", {
            "lecturer": lecturer,
            "subjects": lecturer.subjects.all()
        })

    def post(self, request):
        if not hasattr(request.user, 'lecturer'):
            messages.error(request, "Only lecturers can add questions")
            return redirect("index")

        question_count = request.POST.get("question_count")
        subject_id = request.POST.get("subject")

        if not question_count or not subject_id:
            messages.error(request, "Please select both subject and question count")
            return redirect("select_question_count")

        try:
            question_count = int(question_count)
            if question_count < 1 or question_count > 50:
                messages.error(request, "Please enter a number between 1 and 50")
                return redirect("select_question_count")
        except ValueError:
            messages.error(request, "Please enter a valid number")
            return redirect("select_question_count")

        return redirect(f"{reverse('add_question')}?count={question_count}&subject={subject_id}")

def api_questions_list(request):
    if not request.user.is_authenticated or not hasattr(request.user, 'lecturer'):
        return JsonResponse({'error': 'Unauthorized'}, status=401)

    subject_id = request.GET.get('subject')
    try:
        subject = Subject.objects.get(id=subject_id)
        if subject not in request.user.lecturer.subjects.all():
            return JsonResponse({'error': 'Forbidden'}, status=403)

        questions = TestQuestion.objects.filter(
            subject=subject,
            creator=request.user
        ).values('id', 'question', 'option1', 'option2',
                 'option3', 'option4', 'correct_option',
                 'verified', 'image')

        return JsonResponse({'questions': list(questions)})

    except Subject.DoesNotExist:
        return JsonResponse({'error': 'Առարկան չի գտնվել'}, status=404)

def api_question_detail(request, pk):
    if not request.user.is_authenticated or not hasattr(request.user, 'lecturer'):
        return JsonResponse({'error': 'Unauthorized'}, status=401)

    try:
        question = TestQuestion.objects.get(id=pk)
        if question.creator != request.user:
            return JsonResponse({'error': 'Forbidden'}, status=403)

        data = {
            'id': question.id,
            'text': question.question,
            'option1': question.option1,
            'option2': question.option2,
            'option3': question.option3,
            'option4': question.option4,
            'correct_option': question.correct_option,
            'score': question.score,
            'image': question.image.url if question.image else None
        }
        return JsonResponse(data)

    except TestQuestion.DoesNotExist:
        return JsonResponse({'error': 'Հարցը չի գտնվել'}, status=404)

def get_groups_for_subject(request):
    if not request.user.is_authenticated or not hasattr(request.user, 'lecturer'):
        return JsonResponse({'error': 'Unauthorized'}, status=401)

    subject_id = request.GET.get('subject')
    try:
        subject = Subject.objects.get(id=subject_id)
        # Verify the lecturer teaches this subject
        if subject not in request.user.lecturer.subjects.all():
            return JsonResponse({'error': 'Forbidden'}, status=403)

        groups = Group.objects.filter(major=subject.major)
        data = {
            'groups': [{'id': group.id, 'name': group.name} for group in groups]
        }
        return JsonResponse(data)
    except Subject.DoesNotExist:
        return JsonResponse({'groups': []}, status=404)

@login_required
def results(request):
    if not hasattr(request.user, 'student'):
        return HttpResponseForbidden("Դուք չունեք այս էջը տեսնելու թույլտվություն")

    exam_id = request.GET.get('exam_id')
    student = request.user.student

    if exam_id:
        try:
            try:
                live_exam = LiveStudentExam.objects.get(exam_id=exam_id, student=student)
                result = ExamResult.objects.filter(
                    content_type=ContentType.objects.get_for_model(live_exam),
                    object_id=live_exam.id,
                    student=student
                ).first()
            except LiveStudentExam.DoesNotExist:
                result = ExamResult.objects.filter(
                    content_type=ContentType.objects.get_for_model(MidtermExam),
                    object_id=exam_id,
                    student=student
                ).first()

            if result:
                return redirect('result_detail', result_id=result.id)

            messages.error(request, "Այս քննության համար արդյունքներ չեն գտնվել")
            return redirect('student_dashboard')
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return redirect('student_dashboard')

    # Rest of your existing results view code...
    marks = Mark.objects.filter(user=request.user).select_related('exam')
    exam_results = ExamResult.objects.filter(student=student).select_related('exam')

    # Calculate statistics
    total_tests = marks.count()
    avg_percentage = marks.aggregate(
        avg=models.Avg('got', output_field=models.FloatField()) * 100 / models.F('total')
    )['avg']
    best_score = marks.order_by('-got').first()

    context = {
        'results': marks,
        'total_tests': total_tests,
        'avg_percentage': avg_percentage or 0,
        'best_score': best_score,
    }
    return render(request, 'quiz/results.html', context)

@login_required
def result_detail(request, result_id):
    try:
        # Get the mark
        mark = get_object_or_404(Mark, id=result_id, user=request.user)

        # Find the exam result using GenericRelation
        exam_result = None

        # Try to find through LiveStudentExam first
        try:
            live_exam = LiveStudentExam.objects.get(exam=mark.exam, student=request.user.student)
            exam_result = ExamResult.objects.filter(
                content_type=ContentType.objects.get_for_model(live_exam),
                object_id=live_exam.id,
                student=request.user.student
            ).first()
        except LiveStudentExam.DoesNotExist:
            pass

        # If not found, try through MidtermExam
        if not exam_result:
            exam_result = ExamResult.objects.filter(
                content_type=ContentType.objects.get_for_model(mark.exam),
                object_id=mark.exam.id,
                student=request.user.student
            ).first()

        if not exam_result:
            messages.error(request, "Results not found for this exam")
            return redirect('results')

        # Rest of your view logic...
        points_earned = sum(
            answer.question.score
            for answer in exam_result.answers.all()
            if answer.is_correct
        )
        total_possible_points = sum(
            answer.question.score
            for answer in exam_result.answers.all()
        )
        correct_answers_count = sum(1 for answer in exam_result.answers.all() if answer.is_correct)
        total_questions = exam_result.answers.count()

        context = {
            'result': exam_result,
            'mark': mark,
            'questions': [{
                'question': answer.question,
                'student_answer': answer.answer,
                'correct_answer': answer.question.correct_option,
                'is_correct': answer.is_correct,
                'score': answer.question.score,
                'options': {
                    'A': answer.question.option1,
                    'B': answer.question.option2,
                    'C': answer.question.option3,
                    'D': answer.question.option4,
                }
            } for answer in exam_result.answers.all()],
            'points_earned': points_earned,
            'total_possible_points': total_possible_points,
            'percentage': round((points_earned / total_possible_points) * 100, 2) if total_possible_points > 0 else 0,
            'correct_answers_count': correct_answers_count,
            'total_questions': total_questions,
        }
        return render(request, "quiz/result_detail.html", context)

    except Exception as e:
        messages.error(request, f"Error accessing results: {str(e)}")
        return redirect('results')

# views.py - Example of how to save exam results
def submit_exam(request, exam_id):
    if request.method == 'POST':
        exam = get_object_or_404(MidtermExam, id=exam_id)
        student = request.user.student

        # Calculate score
        score = 0
        student_answers = []

        for question in exam.questions.all():
            answer_key = f'question_{question.id}'
            student_answer = request.POST.get(answer_key)

            if student_answer == question.correct_option:
                score += question.score
                is_correct = True
            else:
                is_correct = False

            student_answer_obj = StudentAnswer.objects.create(
                question=question,
                answer=student_answer,
                is_correct=is_correct
            )
            student_answers.append(student_answer_obj)

        # Create exam result
        exam_result = ExamResult.objects.create(
            exam=exam,
            student=student,
            score=score
        )
        exam_result.answers.set(student_answers)

        # Create mark
        mark = Mark.objects.create(
            total=exam.questions.count(),
            got=score,
            user=request.user,
            exam=exam
        )

        return redirect('results')

def exam_detail_view(request, exam_id):
    exam = get_object_or_404(MidtermExam, id=exam_id)
    return render(request, 'exams/exam_detail.html', {'exam': exam})


from django.contrib import messages
from django.shortcuts import redirect, render
from django.contrib.auth.decorators import login_required
import random

@login_required
def create_exam(request):
    if request.method == 'POST':
        form = MidtermExamForm(request.POST, user=request.user)

        if form.is_valid():
            exam = form.save(commit=False)
            exam.created_by = request.user
            exam.save()

            use_random = form.cleaned_data.get('use_random', False)

            if use_random:
                try:
                    questions_per_student = form.cleaned_data['questions_per_student']
                    subject = form.cleaned_data['subject']
                    group = form.cleaned_data['group']

                    questions = TestQuestion.objects.filter(subject=subject).select_related('category')

                    if not questions.exists():
                        messages.error(request, "Ընթացիկ առարկայի համար հարցեր չկան:")
                        return redirect('lecturer_dashboard')

                    categories = {}
                    uncategorized = []

                    for q in questions:
                        if q.category:
                            if q.category.id not in categories:
                                categories[q.category.id] = {
                                    'name': q.category.name,
                                    'questions': []
                                }
                            categories[q.category.id]['questions'].append(q)
                        else:
                            uncategorized.append(q)

                    if not categories and uncategorized:
                        categories['uncategorized'] = {
                            'name': 'Ընդհանուր',
                            'questions': uncategorized
                        }
                    elif not categories and not uncategorized:
                        messages.error(request, "Առարկայի համար հարցեր չկան:")
                        return redirect('lecturer_dashboard')

                    num_categories = len(categories)
                    if num_categories == 0:
                        messages.error(request, "Հարցերով կատեգորիաներ չեն գտնվել:")
                        return redirect('lecturer_dashboard')

                    base_per_category = questions_per_student // num_categories
                    remainder = questions_per_student % num_categories

                    for cat_id, cat_data in categories.items():
                        available = len(cat_data['questions'])
                        required = base_per_category + (1 if remainder > 0 else 0)
                        if available < required:
                            messages.error(request, f"Կատեգորիան '{cat_data['name']}' ունի միայն {available} հարց, բայց անհրաժեշտ է {required}")
                            return redirect('lecturer_dashboard')

                    students = group.student_set.all()
                    category_ids = list(categories.keys())

                    for student in students:
                        selected_questions = []
                        random.shuffle(category_ids)

                        for i, cat_id in enumerate(category_ids):
                            cat_data = categories[cat_id]
                            take = base_per_category + (1 if i < remainder else 0)
                            selected = random.sample(cat_data['questions'], take)
                            selected_questions.extend(selected)

                        total_score = sum(q.score for q in selected_questions)

                        if total_score != 20:
                            remaining_questions = [q for q in questions if q not in selected_questions]
                            remaining_questions.sort(key=lambda x: x.score)

                            if total_score < 20:
                                for q in remaining_questions:
                                    if total_score + q.score <= 20:
                                        selected_questions.append(q)
                                        total_score += q.score
                                        if total_score == 20:
                                            break
                            else:
                                selected_questions.sort(key=lambda x: x.score)
                                while total_score > 20 and selected_questions:
                                    removed = selected_questions.pop()
                                    total_score -= removed.score

                        if total_score != 20:
                            messages.warning(
                                request,
                                f"{student} ուսանողի համար հնարավոր չեղավ ստանալ ճիշտ 20 միավոր ({total_score})"
                            )

                        live_exam = LiveStudentExam(student=student, exam=exam)
                        live_exam.save()
                        live_exam.questions.set(selected_questions)

                    messages.success(
                        request,
                        f"{students.count()} ուսանողի համար պատահական թեստեր ստեղծվեցին, յուրաքանչյուրում {questions_per_student} հարցով"
                    )
                    return redirect('lecturer_dashboard')

                except Exception as e:
                    messages.error(request, f"Սխալ պատահեց պատահական թեստեր ստեղծելիս: {str(e)}")
                    return redirect('lecturer_dashboard')

            else:
                form.save_m2m()
                students = exam.group.student_set.all()

                for student in students:
                    live_exam = LiveStudentExam(student=student, exam=exam)
                    live_exam.save()
                    live_exam.questions.set(exam.questions.all())

                messages.success(request, "Թեստը հաջողությամբ ստեղծվեց ընտրված հարցերով:")
                return redirect('lecturer_dashboard')

        else:
            messages.error(request, "Սխալ տվյալներ: Խնդրում ենք ստուգել լրացված դաշտերը:")
            return redirect('lecturer_dashboard')

    else:
        messages.error(request, "Թեստի ձևը պետք է ուղարկվի POST մեթոդով:")
        return redirect('lecturer_dashboard')

@login_required
def get_question(request, question_id):
    try:
        question = TestQuestion.objects.get(id=question_id)
        if question.creator != request.user:
            return JsonResponse({'error': 'Permission denied'}, status=403)

        data = {
            'id': question.id,
            'question': question.question,
            'option1': question.option1,
            'option2': question.option2,
            'option3': question.option3,
            'option4': question.option4,
            'correct_option': question.correct_option,
            'score': question.score,
            'image': question.image.url if question.image else None
        }
        return JsonResponse(data)
    except TestQuestion.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)

@login_required
def update_question(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=400)

    try:
        question_id = request.POST.get('id')
        if not question_id:
            return JsonResponse({'error': 'Question ID is required'}, status=400)

        question = TestQuestion.objects.get(id=question_id)
        if question.creator != request.user:
            return JsonResponse({'error': 'Permission denied'}, status=403)

        # Update question fields
        question.question = request.POST.get('question', question.question)
        question.option1 = request.POST.get('option1', question.option1)
        question.option2 = request.POST.get('option2', question.option2)
        question.option3 = request.POST.get('option3', question.option3)
        question.option4 = request.POST.get('option4', question.option4)
        question.correct_option = request.POST.get('correct_option', question.correct_option)
        question.score = int(request.POST.get('score', question.score))

        # Handle image upload
        if 'image' in request.FILES:
            question.image = request.FILES['image']
        elif request.POST.get('remove_image') == 'true':
            question.image.delete()

        question.save()

        return JsonResponse({'success': True})
    except TestQuestion.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)
    except Exception as e:
        import traceback
        print(traceback.format_exc())  # Log full traceback for debugging
        return JsonResponse({'error': str(e)}, status=500)

def feedback_view(request):
    if request.method == 'POST':
        form = FeedbackForm(request.POST)
        if form.is_valid():
            feedback = form.save(commit=False)
            if request.user.is_authenticated:
                feedback.user = request.user
            feedback.save()
            messages.success(request, "Շնորհակալություն ձեր հետադարձ կապի համար")
            return redirect('feedback_success')
    else:
        form = FeedbackForm()

    return render(request, 'feedback/feedback_form.html', {'form': form})

def feedback_success(request):
    return render(request, 'feedback/feedback_success.html')

@method_decorator(login_required, name='dispatch')
class QuestionAPIView(View):
    def get(self, request, pk):
        try:
            question = TestQuestion.objects.get(id=pk)
            if question.creator != request.user:
                return JsonResponse({'error': 'Permission denied'}, status=403)

            data = {
                'id': question.id,
                'question': question.question,
                'option1': question.option1,
                'option2': question.option2,
                'option3': question.option3,
                'option4': question.option4,
                'correct_option': question.correct_option,
                'score': question.score,
                'image': question.image.url if question.image else None
            }
            return JsonResponse(data)

        except TestQuestion.DoesNotExist:
            return JsonResponse({'error': 'Question not found'}, status=404)

    def post(self, request, pk):
        try:
            question = TestQuestion.objects.get(id=pk)
            if question.creator != request.user:
                return JsonResponse({'error': 'Permission denied'}, status=403)

            # Update fields
            question.question = request.POST.get('question', question.question)
            question.option1 = request.POST.get('option1', question.option1)
            question.option2 = request.POST.get('option2', question.option2)
            question.option3 = request.POST.get('option3', question.option3)
            question.option4 = request.POST.get('option4', question.option4)
            question.correct_option = request.POST.get('correct_option', question.correct_option)
            question.score = int(request.POST.get('score', question.score))

            # Handle image upload
            if 'image' in request.FILES:
                question.image = request.FILES['image']

            question.save()

            return JsonResponse({'success': True})

        except TestQuestion.DoesNotExist:
            return JsonResponse({'error': 'Question not found'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

@login_required
def get_question(request, question_id):
    try:
        question = TestQuestion.objects.get(id=question_id)
        if question.creator != request.user:
            return JsonResponse({'error': 'Permission denied'}, status=403)

        data = {
            'id': question.id,
            'question': question.question,
            'option1': question.option1,
            'option2': question.option2,
            'option3': question.option3,
            'option4': question.option4,
            'correct_option': question.correct_option,
            'score': question.score,
            'image': question.image.url if question.image else None
        }
        return JsonResponse(data)
    except TestQuestion.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)

@login_required
def update_question(request):
    if request.method == 'POST':
        try:
            question_id = request.POST.get('id')
            if not question_id:
                return JsonResponse({'error': 'Question ID is required'}, status=400)

            question = TestQuestion.objects.get(id=question_id)
            if question.creator != request.user:
                return JsonResponse({'error': 'Permission denied'}, status=403)

            # Update question fields
            question.question = request.POST.get('question', question.question)
            question.option1 = request.POST.get('option1', question.option1)
            question.option2 = request.POST.get('option2', question.option2)
            question.option3 = request.POST.get('option3', question.option3)
            question.option4 = request.POST.get('option4', question.option4)
            question.correct_option = request.POST.get('correct_option', question.correct_option)
            question.score = int(request.POST.get('score', question.score))

            # Handle image upload if needed
            if 'image' in request.FILES:
                question.image = request.FILES['image']

            question.save()

            return JsonResponse({'success': True})
        except TestQuestion.DoesNotExist:
            return JsonResponse({'error': 'Question not found'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid request method'}, status=400)

def load_questions(request):
    subject_id = request.GET.get('subject_id')
    questions = TestQuestion.objects.filter(subject_id=subject_id).values('id', 'question')
    return JsonResponse(list(questions), safe=False)

def pick_questions_by_dynamic_category(subject, total_score_required, questions_per_test):
    categories = QuestionCategory.objects.filter(subject=subject)
    questions = TestQuestion.objects.filter(subject=subject)

    # Group questions by category
    category_map = {}
    for cat in categories:
        cat_questions = list(questions.filter(category=cat))
        if cat_questions:
            category_map[cat.id] = cat_questions

    if not category_map:
        raise ValueError("No question categories with questions found.")

    num_categories = len(category_map)
    questions_per_category = max(1, questions_per_test // num_categories)

    selected_questions = []
    score_so_far = 0

    for cat_id, question_list in category_map.items():
        if len(question_list) < questions_per_category:
            continue  # Skip if not enough questions in category

        chosen = random.sample(question_list, questions_per_category)
        selected_questions.extend(chosen)
        score_so_far += sum(q.score for q in chosen)

    # Add extra questions to reach total_score_required
    remaining_questions = [q for q in questions if q not in selected_questions]
    random.shuffle(remaining_questions)
    for q in remaining_questions:
        if score_so_far + q.score <= total_score_required:
            selected_questions.append(q)
            score_so_far += q.score
        if score_so_far == total_score_required:
            break

    if score_so_far != total_score_required:
        raise ValueError("Unable to generate balanced questions totaling 20 points.")

    return selected_questions

@login_required
def create_category(request):
    if request.method == 'POST':
        if not hasattr(request.user, 'lecturer'):
            messages.error(request, "Միայն դասախոսները կարող են կատեգորիաներ ստեղծել")
            return redirect('lecturer_dashboard')

        subject_id = request.POST.get('subject')
        name = request.POST.get('name').strip()

        try:
            subject = Subject.objects.get(id=subject_id)
            if subject not in request.user.lecturer.subjects.all():
                messages.error(request, "Դուք չեք դասավանդում այս առարկան")
                return redirect('lecturer_dashboard')

            # Create the category
            category = QuestionCategory.objects.create(
                subject=subject,
                name=name,
                created_by=request.user
            )

            messages.success(request, f"«{name}» կատեգորիան հաջողությամբ ստեղծվել է")
            return redirect('lecturer_dashboard')

        except Subject.DoesNotExist:
            messages.error(request, "Առարկան չի գտնվել")
        except Exception as e:
            messages.error(request, f"Սխալ կատեգորիա ստեղծելիս: {str(e)}")

    return redirect('lecturer_dashboard')

@require_POST
def delete_category(request, category_id):
    try:
        # Make sure category subject is among lecturer's subjects
        category = QuestionCategory.objects.get(
            id=category_id,
            subject__in=request.user.lecturer.subjects.all()
        )
        category_name = category.name
        category.delete()

        return JsonResponse({
            'success': True,
            'message': f'Category "{category_name}" deleted successfully'
        })

    except QuestionCategory.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Category not found or you don’t have permission'
        }, status=404)

    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

